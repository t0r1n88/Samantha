from docx import Document
import pandas as pd
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import time
import re
from openpyxl.styles import Font
from openpyxl.styles import Alignment





# Создаем регулярные выражения
name_table_re = re. compile(r'Наименование государственной услуги:.*?(Реализация\s.+?)[2]')
cat_table_re = re.compile(r'Категории потребителей государственной услуги:.+?(Физические лица,.+?)[3]')

doc = Document('БРИТ.docx')
# doc = Document('БКН.docx')
# doc = Document('БМК.docx')
# последовательность всех таблиц документа
all_tables = doc.tables
print('Всего таблиц в документе:', len(all_tables))

# создаем пустой словарь под данные таблиц
data_tables = {i:None for i in range(len(all_tables))}
# проходимся по таблицам

lst_df =[]
for i, table in enumerate(all_tables):

    # print('\nДанные таблицы №', i)
    # создаем список строк для таблицы `i` (пока пустые)
    data_tables[i] = [[] for _ in range(len(table.rows))]
    # проходимся по строкам таблицы `i`
    for j, row in enumerate(table.rows):
        # проходимся по ячейкам таблицы `i` и строки `j`
        for cell in row.cells:
            # добавляем значение ячейки в соответствующий
            # список, созданного словаря под данные таблиц
            data_tables[i][j].append(cell.text)

    # смотрим извлеченные данные
    # (по строкам) для таблицы `i`
    df = pd.DataFrame(data_tables[i])
    lst_df.append(df)
dct_data = dict()

# Извлекаем
for i in range(1,len(lst_df),3):
    # Получаем совокупный текст таблиц с названием и категорией студентов
    text = lst_df[i].sum().sum()
    # Удаляем символы переноса
    text = text.replace('\n','')

    # Ищем название таблицы
    match_name = re.search(name_table_re,text)
    name_table = match_name.group(1).strip()
    # print(f'Название таблицы {name_table}')
    # Ищем категорию таблицы
    match_cat = re.search(cat_table_re,text)
    name_cat = match_cat.group(1).strip()
    # print(f'Категория таблицы {name_cat}')
    # print('*********************')

    # dct_data[name_table] ={name_cat:''}
    if name_table not in dct_data:
        dct_data[name_table]={name_cat:''}
    else:
        dct_data[name_table].update({name_cat:''})
    if i == 1:
        # удаляем первые 4 строки в датафрейме
        df =  lst_df[3].drop(labels=[0,1,2,3],axis=0)
        # избавляемся от знаков переноса строки

        df = df.applymap(lambda x: x.replace('\n', ''))

    else:
        # удаляем первые 4 строки в датафрейме
        df = lst_df[i+2].drop(labels=[0,1,2,3],axis=0)
        # избавляемся от знаков переноса строки
        df = df.applymap(lambda x:x.replace('\n',''))
    #Добавляем датафрейм в словарь
    dct_data[name_table][name_cat] = df



wb = openpyxl.Workbook()
wb['Sheet']['A1'] = 'Уникальный номер реестровой записи'
wb['Sheet']['B1'] = 'Наименование государственной услуги'
wb['Sheet']['C1'] = 'Категория потребителей государственной услуги'
wb['Sheet']['D1'] = 'Профессии по программам среднего профессионального образования'
wb['Sheet']['E1'] = 'Утверждено в госзадании на год'
wb['Sheet']['F1'] = 'Исполнено на отчетную дату'
wb['Sheet']['G1'] = 'Допустимое отклонение в процентах'
wb['Sheet']['H1'] = 'Допустимое отклонение в ед.'
wb['Sheet']['I1'] = 'отклонение, превышающее допустимое (возможное) значение, %'
wb['Sheet']['J1'] = 'отклонение, превышающее допустимое (возможное) значение, чел.'
wb['Sheet']['K1'] = 'причина отклонения'

wb['Sheet'].column_dimensions['A'].width = 90
wb['Sheet'].column_dimensions['B'].width = 30
wb['Sheet'].column_dimensions['C'].width = 90
wb['Sheet'].column_dimensions['D'].width = 30
wb['Sheet'].column_dimensions['E'].width = 30
wb['Sheet'].column_dimensions['F'].width = 30
wb['Sheet'].column_dimensions['H'].width = 30
wb['Sheet'].column_dimensions['I'].width = 30
wb['Sheet'].column_dimensions['J'].width = 30



for key,value in dct_data.items():
    # Перебираем ключи внутри словаря value
    for cat_key,cat_value in value.items():
        temp_df = cat_value.drop([2,3,4,5,6,7,8],axis=1)

        temp_df.insert(1,'Наименование услуги',key)
        temp_df.insert(2,'Категория потребителей государственной услуги',cat_key)
        for r in dataframe_to_rows(temp_df,index=False,header=False):
            if len(r) != 1:
                wb['Sheet'].append(r)



# Получаем текущее время для того чтобы использовать в названии
t = time.localtime()
current_time = time.strftime('%H_%M_%S', t)
# Сохраняем итоговый файл
wb.save(f'data/Госзадание.xlsx {current_time}.xlsx')